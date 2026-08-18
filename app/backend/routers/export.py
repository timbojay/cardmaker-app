"""
Export router — generates a Word document summarising all cards.
"""

import io
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter
from fastapi.responses import StreamingResponse

from database import get_db

router = APIRouter(prefix="/export", tags=["export"])

PROJECT_ROOT = Path(__file__).parent.parent.parent.parent
PREVIEW_DIR = PROJECT_ROOT / "output" / "previews"


def _get_deck_name(decks_by_id, deck_id):
    return decks_by_id.get(deck_id, deck_id)


@router.get("/cards-doc")
def export_cards_document():
    """Generate a Word document with a table of all cards."""

    # Fetch all decks and cards
    with get_db() as db:
        deck_rows = db.execute("SELECT * FROM decks ORDER BY name").fetchall()
        card_rows = db.execute("SELECT * FROM cards ORDER BY deck_id, title").fetchall()

    decks_by_id = {r["id"]: r["name"] for r in deck_rows}

    # Build the document
    doc = Document()

    # Title
    title = doc.add_heading("Space Junk — Card Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary line
    total = len(card_rows)
    deck_counts = {}
    for row in card_rows:
        dn = _get_deck_name(decks_by_id, row["deck_id"])
        deck_counts[dn] = deck_counts.get(dn, 0) + 1
    summary_parts = [f"{total} cards total"]
    for dn, count in sorted(deck_counts.items()):
        summary_parts.append(f"{dn}: {count}")
    p = doc.add_paragraph(" | ".join(summary_parts))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()  # spacer

    # Group cards by deck
    cards_by_deck = {}
    for row in card_rows:
        deck_name = _get_deck_name(decks_by_id, row["deck_id"])
        if deck_name not in cards_by_deck:
            cards_by_deck[deck_name] = []
        cards_by_deck[deck_name].append(row)

    for deck_name in sorted(cards_by_deck.keys()):
        cards = cards_by_deck[deck_name]

        doc.add_heading(f"{deck_name} ({len(cards)} cards)", level=1)

        # Create table: Title, Benefits, Costs, Description
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ["Title", "Benefits", "Costs", "Description"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for card in cards:
            row_cells = table.add_row().cells

            # Title
            row_cells[0].text = card["title"]
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

            # Benefits
            benefits = json.loads(card["benefits"]) if card["benefits"] else {}
            benefit_parts = []
            for icon, count in benefits.items():
                if count > 0:
                    benefit_parts.append(f"{icon} x{count}")
            row_cells[1].text = ", ".join(benefit_parts) if benefit_parts else "—"
            for paragraph in row_cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

            # Costs
            costs = json.loads(card["costs"]) if card["costs"] else {}
            cost_parts = []
            for icon, count in costs.items():
                if count > 0:
                    cost_parts.append(f"{icon} x{count}")
            row_cells[2].text = ", ".join(cost_parts) if cost_parts else "—"
            for paragraph in row_cells[2].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

            # Description
            row_cells[3].text = card["description"] or ""
            for paragraph in row_cells[3].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.3)
            row.cells[1].width = Inches(1.8)
            row.cells[2].width = Inches(1.8)
            row.cells[3].width = Inches(2.6)

        doc.add_paragraph()  # spacer between decks

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=space_junk_cards.docx"},
    )
